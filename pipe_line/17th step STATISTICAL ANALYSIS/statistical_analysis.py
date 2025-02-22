import os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# Specify the folder paths
input_folder = "15th step ECDF PRUNNING/pruned_dataset/" # After we did the not imputed values statistic analysis we did the imputed values statistic analysis ("16th step KNN IMPUTATION/imputed_dataset/")
output_folder = "17th step STATISTICAL ANALYSIS/pruned_dataset/" # We will save the statistical analysis of imputed_dataset to "17th step STATISTICAL ANALYSIS/imputed_dataset/"

# Function to set the cell margins in a Word table
def set_cell_margins(cell, **kwargs):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for k, v in kwargs.items():
        node = OxmlElement(f'w:{k}')
        node.set(qn('w:w'), str(v))
        tcMar.append(node)
    tcPr.append(tcMar)

# Function to generate the ECDF plot for a DataFrame
def generate_ecdf_plot(df, output_folder, file_name):
    numeric_cols = df.select_dtypes(include='number').columns
    
    # Ensure the output folder exists
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # Generate ECDF plots for each numeric column
    for col in numeric_cols:
        plt.figure(figsize=(28, 18))  
        
        x = np.sort(df[col].dropna())  
        y = np.arange(1, len(x)+1) / len(x)
        
        plt.plot(x, y, marker='o', linestyle='-', color='blue', label=f'ECDF of {col}')
        plt.fill_between(x, y, color='lightblue', alpha=0.5)
        
        plt.xlabel(f'{col} Values', fontsize=16)
        plt.ylabel('ECDF', fontsize=16)
        plt.title(f'ECDF of {col}', fontsize=20, fontweight='bold')
        plt.grid(True)
        plt.xticks(fontsize=14)
        plt.yticks(fontsize=14)
        plt.legend(fontsize=14)
        
        # Save the ECDF plot as an image
        ecdf_image_path = os.path.join(output_folder, f'{file_name}_{col}_ecdf.png')
        
        # Ensure the subdirectory for the specific file exists
        ecdf_dir = os.path.dirname(ecdf_image_path)
        if not os.path.exists(ecdf_dir):
            os.makedirs(ecdf_dir)
        
        plt.savefig(ecdf_image_path, bbox_inches='tight')
        plt.close()

        # Return the path for adding it to the Word document
        yield ecdf_image_path

# Function to process the Excel file
def process_excel_file(file_path, output_folder):
    df = pd.read_excel(file_path)
    # Drop rows with all missing values
    df.dropna(how='all', inplace=True)

    # Filter out columns with zero variance and dummy columns
    numeric_df = df.select_dtypes(include='number')
    dummy_columns = ['Year', 'Month', 'Day', 'COVID-19']
    non_dummy_cols = numeric_df.columns[numeric_df.var(numeric_only=True) != 0]
    non_dummy_df = numeric_df[non_dummy_cols]

    non_dummy_cols = [
        col for col in non_dummy_df.columns
        if col not in dummy_columns and not (non_dummy_df[col].isin([0, 1]).all() and non_dummy_df[col].nunique() <= 2)
    ]
    filtered_df = non_dummy_df[non_dummy_cols]

    # Generate the correlation matrix heatmap
    correlation_matrix = filtered_df.corr()

    plt.figure(figsize=(30, 20))
    sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', vmin=-1, vmax=1)
    heatmap_image_path = os.path.join(output_folder, os.path.basename(file_path).replace('.xlsx', '_heatmap.png'))
    plt.savefig(heatmap_image_path)
    plt.close()

    # Generate the statistical summary
    stats_summary = filtered_df.describe(percentiles=[.2, .4, .6, .8]).T
    stats_summary['Mean'] = stats_summary['mean']
    stats_summary['Standard Deviation'] = stats_summary['std']
    stats_summary['Minimum Value'] = stats_summary['min']
    stats_summary['20th Percentile (Q1)'] = stats_summary['20%']
    stats_summary['40th Percentile (Q2)'] = stats_summary['40%']
    stats_summary['Median (Q3)'] = stats_summary['50%']
    stats_summary['60th Percentile (Q4)'] = stats_summary['60%']
    stats_summary['80th Percentile (Q5)'] = stats_summary['80%']
    stats_summary['Maximum Value'] = stats_summary['max']
    
    stats_summary = stats_summary.drop(columns=['mean', 'std', 'min', '20%', '40%', '50%', '60%', '80%', 'max'])
    
    # Calculate Missing Data and Valid Data as percentages
    stats_summary['Missing Data'] = (filtered_df.isnull().sum() / len(filtered_df)) * 100
    stats_summary['Valid Data'] = (filtered_df.notnull().sum() / len(filtered_df)) * 100
    stats_summary['Total Observations'] = len(filtered_df)

    
    stats_summary = stats_summary[
        ['Total Observations', 'Valid Data', 'Missing Data',
         'Mean', 'Standard Deviation', 'Minimum Value', 
         '20th Percentile (Q1)', '40th Percentile (Q2)', 'Median (Q3)', 
         '60th Percentile (Q4)', '80th Percentile (Q5)', 'Maximum Value']
    ].round(2)

    # Create a Word document to save the analysis
    doc = Document()
    section = doc.sections[0]
    section.page_height = Pt(800)
    section.page_width = Pt(1000)
    section.top_margin = Pt(40)
    section.bottom_margin = Pt(40)
    section.left_margin = Pt(40)
    section.right_margin = Pt(40)

    # Add the title and the correlation matrix heatmap
    doc.add_heading('Statistical Analysis', level=1)
    doc.add_heading('Correlation Matrix Heatmap', level=2)
    doc.add_picture(heatmap_image_path, width=Inches(14), height=Inches(9))
    doc.add_heading('Fundamental Statistical Analysis', level=2)

    # Add the statistical summary table
    table = doc.add_table(rows=1, cols=stats_summary.shape[1] + 1)
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    for j, column in enumerate(stats_summary.columns):
        hdr_cells[j + 1].text = column

    # Bold the header row
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs
            if run:
                run[0].bold = True
    # Add the data rows
    for i in range(stats_summary.shape[0]):
        if stats_summary.iloc[i].isna().all() or stats_summary.iloc[i]['Valid Data'] == 0:
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = stats_summary.index[i]
        for j in range(stats_summary.shape[1]):
            row_cells[j + 1].text = str(stats_summary.iloc[i, j])
    # Format the table
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.2)
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            set_cell_margins(cell, top=0, left=0, bottom=0, right=0)
    
    # Highlight the header row
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if i == 0:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')))
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Center the table
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the ECDF plots
    doc.add_heading('Empirical Cumulative Distribution Function (ECDF) Plots', level=2)
    for ecdf_image_path in generate_ecdf_plot(filtered_df, output_folder, os.path.basename(file_path).replace('.xlsx', '')):
        doc.add_picture(ecdf_image_path, width=Inches(12), height=Inches(8))  # ECDF image size doubled
        os.remove(ecdf_image_path)  # Remove the ECDF image after it's added

    # Save the Word document
    output_file_path = os.path.join(output_folder, os.path.basename(file_path).replace('.xlsx', '_analysis.docx'))
    doc.save(output_file_path)
    os.remove(heatmap_image_path)  # Remove heatmap after adding to Word

    print(f'Statistical analysis and ECDF graphs saved to {output_file_path}')

# Function to process all files in a folder
def process_all_files_in_folder(input_folder, output_folder):
    for file_name in os.listdir(input_folder):
        if file_name.endswith('.xlsx'):
            file_path = os.path.join(input_folder, file_name)
            process_excel_file(file_path, output_folder)

# Ensure the output folder exists and process all files in the input folder
os.makedirs(output_folder, exist_ok=True)
process_all_files_in_folder(input_folder, output_folder)
